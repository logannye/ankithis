"""Parse DOCX files using python-docx. Uses heading styles for section detection."""

from __future__ import annotations

from dataclasses import dataclass, field

from docx import Document


@dataclass
class ParsedSection:
    title: str | None
    level: int
    paragraphs: list[str] = field(default_factory=list)


def parse_docx(file_path: str) -> list[ParsedSection]:
    """Parse a DOCX file into sections based on heading styles."""
    doc = Document(file_path)
    sections: list[ParsedSection] = []
    current_section = ParsedSection(title=None, level=1)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()

        if style_name.startswith("heading"):
            # Save current section
            if current_section.paragraphs or current_section.title:
                sections.append(current_section)
            # Extract level from "Heading 1", "Heading 2", etc.
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            current_section = ParsedSection(title=text, level=level)
        elif style_name == "title":
            if current_section.paragraphs or current_section.title:
                sections.append(current_section)
            current_section = ParsedSection(title=text, level=1)
        else:
            current_section.paragraphs.append(text)

    # Final section
    if current_section.paragraphs or current_section.title:
        sections.append(current_section)

    if not sections:
        sections = [ParsedSection(title=None, level=1)]

    return sections
